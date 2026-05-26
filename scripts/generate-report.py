#!/usr/bin/env python3
"""
Generador de Informes OSINT a partir de la Plantilla VIP de Inteligencia Digital.
Lee datos JSON desde stdin, rellena la plantilla DOCX y guarda el resultado.
Compatible con Plantilla_de_Informes_VIP.docx
"""

import sys
import json
import hashlib
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Dynamic paths - works both locally and on Railway.app
APP_ROOT = os.environ.get("APP_ROOT", os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
TEMPLATE_PATH = os.path.join(APP_ROOT, "upload", "Plantilla_de_Informes_VIP.docx")
OUTPUT_DIR = os.path.join(APP_ROOT, "download", "reports")

def generate_report(data: dict) -> str:
    """Generate a DOCX report from the VIP template filled with scan data."""
    
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    
    doc = Document(TEMPLATE_PATH)
    
    # ── Extract data ──
    scan = data.get("scan", {})
    results = data.get("results", [])
    full_name = scan.get("fullName", "")
    cedula = scan.get("cedula", "") or ""
    email = scan.get("email", "") or ""
    phone = scan.get("phone", "") or ""
    scan_id = scan.get("id", "N/A")
    created_at = scan.get("createdAt", "")
    
    today = datetime.now().strftime("%d/%m/%Y")
    now_iso = datetime.now().strftime("%d/%m/%Y %H:%M")
    
    # Categorize results
    credential_breaches = [r for r in results if r.get("category") == "credential_breach"]
    password_exposure = [r for r in results if r.get("category") == "password_exposure"]
    personal_exposure = [r for r in results if r.get("category") == "personal_exposure"]
    social_media = [r for r in results if r.get("category") == "social_media"]
    data_broker = [r for r in results if r.get("category") == "data_broker"]
    dark_web = [r for r in results if r.get("category") in ("dark_web_mention", "paste_site")]
    document_exposure = [r for r in results if r.get("category") == "document_exposure"]
    
    # Calculate risk
    critical_count = len([r for r in results if r.get("severity") == "critical"])
    high_count = len([r for r in results if r.get("severity") == "high"])
    medium_count = len([r for r in results if r.get("severity") == "medium"])
    low_count = len([r for r in results if r.get("severity") == "low"])
    info_count = len([r for r in results if r.get("severity") == "info"])
    
    risk_score = min(100, critical_count * 30 + high_count * 15 + medium_count * 5 + low_count * 2)
    if risk_score >= 70:
        risk_level = "CRITICO"
    elif risk_score >= 40:
        risk_level = "ALTO"
    elif risk_score >= 15:
        risk_level = "MODERADO"
    else:
        risk_level = "BAJO"
    
    # ── Table 0: Metadata (4 rows x 4 cols) ──
    table0 = doc.tables[0]
    table0.rows[0].cells[1].text = f"OSINT-{scan_id[:8].upper()}"
    table0.rows[0].cells[3].text = today
    table0.rows[1].cells[1].text = "OSINT Scanner Automatizado"
    table0.rows[1].cells[3].text = "1.0"
    table0.rows[2].cells[1].text = "Sistema Automatizado"
    table0.rows[2].cells[3].text = "FINAL"
    table0.rows[3].cells[1].text = full_name
    
    # ── Table 1: Aviso Legal - Keep original, no modification needed ──
    
    # ── Table 2: Hallazgo Principal (1 row x 1 col) ──
    table2 = doc.tables[2]
    finding_text = f"HALLAZGO PRINCIPAL\n"
    if results:
        finding_text += f"Se identificaron {len(results)} hallazgos para el sujeto {full_name} mediante analisis de fuentes abiertas (OSINT).\n\n"
        if credential_breaches:
            finding_text += f"- {len(credential_breaches)} filtraciones de credenciales detectadas en bases de datos comprometidas.\n"
        if password_exposure:
            finding_text += f"- {len(password_exposure)} exposiciones de contraseñas en filtraciones conocidas.\n"
        if personal_exposure:
            finding_text += f"- {len(personal_exposure)} exposiciones de datos personales en sitios web publicos.\n"
        if social_media:
            finding_text += f"- {len(social_media)} perfiles en redes sociales identificados y verificados.\n"
        if data_broker:
            finding_text += f"- {len(data_broker)} registros en brokers de datos y directorios publicos.\n"
        if dark_web:
            finding_text += f"- {len(dark_web)} menciones en filtraciones y sitios de paste/dark web.\n"
        if document_exposure:
            finding_text += f"- {len(document_exposure)} documentos expuestos con datos del sujeto.\n"
    else:
        finding_text += f"No se identificaron hallazgos significativos para los datos proporcionados del sujeto {full_name}. La huella digital presenta un nivel de exposicion bajo segun las fuentes consultadas.\n"
    
    finding_text += f"\nNIVEL DE RIESGO / RELEVANCIA\n"
    finding_text += f"Nivel: {risk_level} (Puntuacion: {risk_score}/100)\n"
    finding_text += f"Hallazgos - Criticos: {critical_count} | Altos: {high_count} | Medios: {medium_count} | Bajos: {low_count} | Info: {info_count}"
    
    table2.rows[0].cells[0].text = finding_text
    
    # ── Table 3: Identidad del Sujeto (10 rows x 2 cols) ──
    table3 = doc.tables[3]
    table3.rows[0].cells[1].text = full_name
    table3.rows[1].cells[1].text = cedula if cedula else "No proporcionado"
    table3.rows[2].cells[1].text = ""  # Alias
    table3.rows[3].cells[1].text = ""  # Fecha nacimiento
    table3.rows[4].cells[1].text = "Colombia"
    table3.rows[5].cells[1].text = ""  # Ciudad
    table3.rows[6].cells[1].text = ""  # Ocupacion
    table3.rows[7].cells[1].text = ""  # Empleador
    table3.rows[8].cells[1].text = phone if phone else "No proporcionado"
    table3.rows[9].cells[1].text = email if email else "No proporcionado"
    
    # ── Table 4: Fotografias e imagenes (4 rows x 5 cols) ──
    table4 = doc.tables[4]
    img_results = [r for r in results if "foto" in r.get("title", "").lower() 
                   or "imagen" in r.get("title", "").lower() 
                   or "profile" in r.get("title", "").lower()
                   or "social_media" == r.get("category", "")]
    for i, r in enumerate(img_results[:3]):
        if i + 1 < len(table4.rows):
            table4.rows[i + 1].cells[0].text = str(i + 1)
            table4.rows[i + 1].cells[1].text = r.get("url", "")[:80]
            table4.rows[i + 1].cells[2].text = today
            table4.rows[i + 1].cells[3].text = "-"
            table4.rows[i + 1].cells[4].text = r.get("description", "")[:100]
    
    # ── Table 5: Redes Sociales (9 rows x 5 cols) ──
    table5 = doc.tables[5]
    platform_map = {
        "linkedin": 1, "twitter": 2, "x": 2, "instagram": 3,
        "facebook": 4, "github": 5, "tiktok": 6, "telegram": 7
    }
    for r in social_media:
        url = r.get("url", "").lower()
        title = r.get("title", "").lower()
        for platform, row_idx in platform_map.items():
            if platform in url or platform in title:
                if row_idx < len(table5.rows):
                    table5.rows[row_idx].cells[1].text = r.get("url", "")[:60]
                    table5.rows[row_idx].cells[2].text = "Identificado"
                    table5.rows[row_idx].cells[3].text = today
                    severity = r.get("severity", "info")
                    table5.rows[row_idx].cells[4].text = "Alto" if severity in ("critical", "high") else "Medio" if severity == "medium" else "Bajo"
    
    # ── Table 6: Correos y Brechas (4 rows x 4 cols) ──
    table6 = doc.tables[6]
    breach_results = credential_breaches + password_exposure + dark_web
    for i, r in enumerate(breach_results[:3]):
        if i + 1 < len(table6.rows):
            table6.rows[i + 1].cells[0].text = email or r.get("dataFound", "")
            table6.rows[i + 1].cells[1].text = r.get("source", "")
            has_breach = r.get("category") in ("credential_breach", "password_exposure")
            breach_count = len(credential_breaches)
            table6.rows[i + 1].cells[2].text = f"SI - {breach_count} brecha(s)" if has_breach else "NO"
            table6.rows[i + 1].cells[3].text = r.get("description", "")[:80] or "PII / combinacion"
    
    # ── Table 7: Dominios e Infraestructura (3 rows x 5 cols) ──
    table7 = doc.tables[7]
    domain_results = [r for r in results if "dominio" in r.get("title", "").lower() 
                      or "whois" in r.get("source", "").lower() 
                      or "domain" in r.get("category", "")
                      or "infraestructura" in r.get("title", "").lower()]
    for i, r in enumerate(domain_results[:2]):
        if i + 1 < len(table7.rows):
            table7.rows[i + 1].cells[0].text = r.get("url", "")[:60] or r.get("dataFound", "")
            table7.rows[i + 1].cells[1].text = "Ver WHOIS"
            table7.rows[i + 1].cells[2].text = today
            table7.rows[i + 1].cells[3].text = "Por verificar"
            table7.rows[i + 1].cells[4].text = "Por determinar"
    
    # ── Table 8: Red de Relaciones (6 rows x 4 cols) ──
    table8 = doc.tables[8]
    rel_results = [r for r in results if "relacion" in r.get("title", "").lower() 
                   or "vinculo" in r.get("title", "").lower()
                   or "asociado" in r.get("title", "").lower()]
    for i, r in enumerate(rel_results[:5]):
        if i + 1 < len(table8.rows):
            table8.rows[i + 1].cells[0].text = r.get("dataFound", "")[:50]
            table8.rows[i + 1].cells[1].text = "Por clasificar"
            table8.rows[i + 1].cells[2].text = r.get("source", "")
            table8.rows[i + 1].cells[3].text = r.get("description", "")[:80] or "-"
    
    # ── Table 9: Vinculos Empresariales RUES (4 rows x 5 cols) ──
    table9 = doc.tables[9]
    corp_results = [r for r in results if "empresa" in r.get("title", "").lower() 
                    or "rues" in r.get("source", "").lower() 
                    or "corporate" in r.get("category", "")
                    or "sociedad" in r.get("title", "").lower()]
    for i, r in enumerate(corp_results[:3]):
        if i + 1 < len(table9.rows):
            table9.rows[i + 1].cells[0].text = r.get("dataFound", "")[:50]
            table9.rows[i + 1].cells[1].text = "N/A"
            table9.rows[i + 1].cells[2].text = "Por verificar"
            table9.rows[i + 1].cells[3].text = "Verificar"
            table9.rows[i + 1].cells[4].text = r.get("source", "")
    
    # ── Table 10: Antecedentes Judiciales (8 rows x 4 cols) ──
    table10 = doc.tables[10]
    judicial_results = [r for r in results if "judicial" in r.get("category", "") 
                        or "antecedente" in r.get("title", "").lower() 
                        or "ofac" in r.get("source", "").lower()
                        or "sancion" in r.get("title", "").lower()
                        or "lista" in r.get("title", "").lower()]
    for i in range(1, len(table10.rows)):
        matched = None
        for r in judicial_results:
            source_lower = r.get("source", "").lower()
            cell_text = table10.rows[i].cells[0].text.lower()
            if any(keyword in source_lower for keyword in cell_text.split()):
                matched = r
                break
        if matched:
            table10.rows[i].cells[1].text = matched.get("description", "Ver detalle")[:80]
            table10.rows[i].cells[2].text = today
            table10.rows[i].cells[3].text = matched.get("severity", "info")
        else:
            table10.rows[i].cells[1].text = "Sin resultados"
            table10.rows[i].cells[2].text = today
            table10.rows[i].cells[3].text = "-"
    
    # ── Table 11: Presencia en Medios (4 rows x 4 cols) ──
    table11 = doc.tables[11]
    media_results = [r for r in results if "medio" in r.get("title", "").lower() 
                     or "prensa" in r.get("category", "") 
                     or "noticia" in r.get("title", "").lower()
                     or "articulo" in r.get("title", "").lower()]
    for i, r in enumerate(media_results[:3]):
        if i + 1 < len(table11.rows):
            table11.rows[i + 1].cells[0].text = r.get("source", "")
            table11.rows[i + 1].cells[1].text = r.get("title", "")[:80]
            table11.rows[i + 1].cells[2].text = today
            table11.rows[i + 1].cells[3].text = r.get("url", "")[:80]
    
    # ── Table 12: Linea de Tiempo (5 rows x 4 cols) ──
    table12 = doc.tables[12]
    all_important = [r for r in results if r.get("severity") in ("critical", "high", "medium")]
    for i, r in enumerate(all_important[:4]):
        if i + 1 < len(table12.rows):
            table12.rows[i + 1].cells[0].text = today
            table12.rows[i + 1].cells[1].text = r.get("title", "")[:100]
            table12.rows[i + 1].cells[2].text = r.get("source", "")
            sev = r.get("severity", "info")
            table12.rows[i + 1].cells[3].text = "Alta" if sev in ("critical", "high") else "Media" if sev == "medium" else "Baja"
    
    # ── Table 13: Indicadores de Riesgo IoR (4 rows x 5 cols) ──
    table13 = doc.tables[13]
    risk_indicators = [r for r in results if r.get("severity") in ("critical", "high")]
    for i, r in enumerate(risk_indicators[:3]):
        if i + 1 < len(table13.rows):
            table13.rows[i + 1].cells[0].text = str(i + 1)
            table13.rows[i + 1].cells[1].text = r.get("title", "")[:100]
            table13.rows[i + 1].cells[2].text = r.get("url", "")[:60] or r.get("source", "")
            sev = r.get("severity", "info")
            table13.rows[i + 1].cells[3].text = "ALTA" if sev == "critical" else "MEDIA"
            cat = r.get("category", "")
            rec_map = {
                "credential_breach": "Cambiar contraseñas inmediatamente. Habilitar 2FA en todas las cuentas asociadas al correo comprometido.",
                "password_exposure": "Rotar todas las contraseñas asociadas. Implementar gestor de contraseñas y autenticacion multi-factor.",
                "personal_exposure": "Solicitar eliminacion de datos al sitio. Revisar y restringir configuracion de privacidad en todas las plataformas.",
                "dark_web_mention": "Monitoreo continuo de actividad sospechosa. Considerar alertas de fraude y bloqueo preventivo.",
                "paste_site": "Cambiar credenciales comprometidas inmediatamente. Verificar actividad sospechosa en cuentas financieras.",
                "data_broker": "Ejercer derecho de supresion bajo Ley 1581 de 2012 (Habeas Data). Contactar directamente al broker de datos.",
                "social_media": "Revisar configuracion de privacidad en perfiles identificados. Limitar informacion publica disponible.",
                "document_exposure": "Solicitar eliminacion del documento expuesto. Verificar alcance completo de la exposicion y posibles usos indebidos.",
            }
            table13.rows[i + 1].cells[4].text = rec_map.get(cat, "Investigar y tomar medidas segun el tipo de exposicion detectada.")
    
    # ── Table 14: Conclusiones del Analista (1 row x 1 col) ──
    table14 = doc.tables[14]
    conclusions = f"CONCLUSIONES DEL ANALISIS DE INTELIGENCIA DIGITAL\n\n"
    conclusions += f"1. Confirmacion de hallazgos: Se identificaron {len(results)} resultados relevantes para los datos proporcionados del sujeto {full_name}"
    if cedula:
        conclusions += f", documento {cedula}"
    conclusions += f". "
    if results:
        conclusions += f"La investigacion OSINT confirma la existencia de exposicion de datos personales en multiples fuentes abiertas.\n\n"
    else:
        conclusions += f"La investigacion OSINT no revelo hallazgos significativos en las fuentes consultadas.\n\n"
    
    if critical_count > 0:
        conclusions += f"2. Nivel de certeza: ALTO - Se detectaron {critical_count} hallazgos criticos que requieren atencion inmediata, incluyendo filtraciones de credenciales y exposicion de datos sensibles. La confiabilidad de estas fuentes es alta al tratarse de bases de datos de brechas verificadas.\n\n"
    elif high_count > 0:
        conclusions += f"2. Nivel de certeza: MEDIO-ALTO - Se detectaron {high_count} hallazgos de severidad alta que indican exposicion significativa de datos personales en fuentes publicas y directorios.\n\n"
    else:
        conclusions += f"2. Nivel de certeza: MEDIO - Los hallazgos detectados son de severidad moderada o baja, lo que sugiere una exposicion limitada segun las fuentes consultadas. No se descarta la existencia de datos en fuentes no cubiertas por este analisis.\n\n"
    
    conclusions += f"3. Gaps de informacion: No fue posible verificar todas las fuentes de datos publicos de forma automatizada. Se recomienda verificacion manual en las siguientes fuentes: RUES (Registro Unico Empresarial), Rama Judicial, listas restrictivas internacionales (OFAC, ONU), y bases de datos de propiedad intelectual.\n\n"
    
    conclusions += f"4. Recomendaciones de seguimiento: "
    if critical_count > 0:
        conclusions += "ACCION INMEDIATA REQUERIDA - Cambiar todas las credenciales comprometidas, habilitar autenticacion de dos factores en cada cuenta, solicitar la eliminacion de datos personales en los sitios identificados y considerar el monitoreo continuo de identidad. Escalar al area de seguridad informatica."
    elif high_count > 0:
        conclusions += "Se recomienda revisar la configuracion de privacidad en las plataformas identificadas, cambiar contraseñas asociadas al correo, monitorear la actividad de las cuentas vinculadas y ejercer los derechos de Habeas Data ante los brokers de datos detectados."
    else:
        conclusions += "Se recomienda mantener practicas de higiene digital, verificar periodicamente la exposicion de datos personales, considerar el uso de servicios de monitoreo de identidad y mantener actualizadas las configuraciones de privacidad en redes sociales."
    
    table14.rows[0].cells[0].text = conclusions
    
    # ── Table 15: Cadena de Evidencia (7 rows x 7 cols) ──
    table15 = doc.tables[15]
    all_sources = {}
    for r in results:
        src = r.get("source", "")
        if src not in all_sources:
            all_sources[src] = []
        all_sources[src].append(r)
    
    row_idx = 1
    fid = 1
    for source, source_results in all_sources.items():
        for r in source_results[:2]:
            if row_idx < len(table15.rows):
                table15.rows[row_idx].cells[0].text = f"F{fid}"
                table15.rows[row_idx].cells[1].text = source
                table15.rows[row_idx].cells[2].text = r.get("url", "")[:60] or "Resultado de busqueda web"
                table15.rows[row_idx].cells[3].text = now_iso
                table15.rows[row_idx].cells[4].text = "NO"
                # Generate hash from URL
                url_bytes = r.get("url", "").encode() or r.get("title", "").encode()
                table15.rows[row_idx].cells[5].text = f"SHA256: {hashlib.sha256(url_bytes).hexdigest()[:16]}..."
                sev = r.get("severity", "info")
                table15.rows[row_idx].cells[6].text = "A" if sev in ("critical", "high") else "M" if sev == "medium" else "B"
                row_idx += 1
                fid += 1
    
    # ── Table 16: Firma y Aprobacion (2 rows x 3 cols) ──
    table16 = doc.tables[16]
    table16.rows[1].cells[0].text = f"Firma\n\nOSINT Scanner Automatizado\nFecha: {today}"
    
    # ── Update executive summary paragraph ──
    for p in doc.paragraphs:
        if "XXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXX" in p.text:
            p.text = f"Con base en la investigación OSINT realizada al sujeto {full_name}, identificando documento {cedula or 'N/A'}, correo {email or 'N/A'} y teléfono {phone or 'N/A'}, se ejecutaron búsquedas en múltiples motores de fuentes abiertas. El análisis reveló {len(results)} hallazgos con un nivel de riesgo {risk_level} (puntuación {risk_score}/100). {'Se detectaron filtraciones de credenciales y exposición de datos personales que requieren atención inmediata.' if critical_count > 0 else 'No se detectaron hallazgos críticos, pero se recomienda monitoreo continuo.' if len(results) == 0 else 'Los hallazgos detectados indican un nivel de exposición que requiere medidas de protección.'}"
            for run in p.runs:
                run.font.size = Pt(10)
            break
    
    # ── Save ──
    safe_name = full_name.replace(" ", "_").replace("/", "_").replace("\\", "_")
    filename = f"Informe_OSINT_{safe_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    output_path = os.path.join(OUTPUT_DIR, filename)
    doc.save(output_path)
    
    return json.dumps({
        "success": True,
        "filePath": output_path,
        "fileName": filename
    })


if __name__ == "__main__":
    try:
        input_data = json.loads(sys.stdin.read())
        result = generate_report(input_data)
        print(result)
    except Exception as e:
        print(json.dumps({
            "success": False,
            "error": str(e)
        }))
        sys.exit(1)
